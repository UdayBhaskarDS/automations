# import sys
# from docx import Document
# import re

# file_path = sys.argv[1]  # Get path from command argument

# def clean_text(text: str) -> str:
#     text = re.sub(r'\r?\n|\r', ' ', text)
#     text = re.sub(r'[^a-zA-Z0-9\s\.\,\?\!\-\:\%]', '', text)
#     text = re.sub(r'\s+', ' ', text)
#     return text.strip()

# doc = Document(file_path)
# paragraphs = [clean_text(p.text) for p in doc.paragraphs if p.text.strip()]
# output = " ".join(paragraphs)

# print(output)

# import sys
# import json
# import docx

# # Get DOCX file path from n8n argument
# docx_path = sys.argv[1]

# # Open and read DOCX
# doc = docx.Document(docx_path)
# full_text = []
# for para in doc.paragraphs:
#     full_text.append(para.text)

# # Combine into single string
# text_content = "\n".join(full_text)

# # Output JSON for n8n
# print(json.dumps({"text": text_content}))


#print("This script is designed to convert a DOCX file to plain text. Please provide the path to the DOCX file as a command line argument.")
### working for local testing
import sys
import json
from docx import Document

# Input path from n8n
docx_path = sys.argv[1]

# Load the docx file
doc = Document(docx_path)

# Extract all text
full_text = "\n".join([para.text for para in doc.paragraphs])

# Output JSON with same schema as PDF branch
print(json.dumps({"text": full_text}))

# import sys
# import json
# from docx import Document

# print("test sys.argv:", sys.argv[0])
# docx_path = sys.argv[0]  # full local path from n8n

# doc = Document(docx_path)

# full_text = "\n".join([p.text for p in doc.paragraphs])
# print(json.dumps({"text": full_text}))


